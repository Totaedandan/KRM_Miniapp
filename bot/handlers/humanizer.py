"""
Хэндлер Хуманайзера (StealthGPT).

Логика:
  - Пользователь нажимает "✍️ Хуманайзер"
  - Видит меню с настройками (тон, режим, бизнес-мод) и балансом токенов
  - Нажимает "Хуманизировать текст" → отправляет текст или файл (.txt/.docx/.doc)
  - Бот считает стоимость (0.5 токена/слово, ×10 в бизнес-моде)
  - Если токенов хватает — хуманизирует и отправляет результат
  - Если не хватает — предлагает купить токены

Пользователи в whitelist — хуманайзер бесплатно (без списания токенов).
"""

import math
import logging
import traceback

from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, Document, BufferedInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

from database import db
from keyboards.humanizer_kb import humanizer_menu_kb, tone_kb, mode_kb
from keyboards.main_kb import main_menu_kb, cancel_kb
from keyboards.payment_kb import token_packages_kb
from services.humanizer_service import (
    humanize_text, calculate_humanizer_cost, TONE_OPTIONS, MODE_OPTIONS
)
from config import settings

logger = logging.getLogger(__name__)
router = Router()

MENU_BUTTONS = {"❌ Отмена", "🤖 ИИ-чат", "✍️ Хуманайзер", "👤 Профиль", "💳 Купить токены", "📄 Turnitin"}


class HumanizerFSM(StatesGroup):
    waiting_text = State()


# ── Главное меню Хуманайзера ──────────────────────────────────────────────────

async def _humanizer_text(tg_id: int, tone: str, mode: str, business: bool) -> str:
    balance = await db.get_token_balance(tg_id)
    tone_label = TONE_OPTIONS.get(tone, tone)
    mode_label = MODE_OPTIONS.get(mode, mode)
    cost_label = "5 🪙/слово" if business else "0.5 🪙/слово"
    return (
        "✍️ <b>Хуманайзер</b>\n\n"
        "Делаю ИИ-тексты неотличимыми от человеческих.\n\n"
        f"⚙️ <b>Настройки:</b>\n"
        f"• Тон: <b>{tone_label}</b>\n"
        f"• Режим: <b>{mode_label}</b>\n"
        f"• Бизнес-мод: <b>{'✅ Вкл' if business else '❌ Выкл'}</b>\n\n"
        f"💰 Стоимость: <b>{cost_label}</b>\n"
        f"🪙 Ваш баланс: <b>{balance:.1f} токенов</b>\n\n"
        "📎 Отправьте текст сообщением или файлом <b>.txt / .docx / .doc</b>"
    )


@router.message(F.text == "✍️ Хуманайзер")
async def humanizer_menu(message: Message, state: FSMContext):
    await state.clear()
    data = await state.get_data()
    tone     = data.get("tone", "College")
    mode     = data.get("mode", "Medium")
    business = data.get("business", False)
    text = await _humanizer_text(message.from_user.id, tone, mode, business)
    await message.answer(text, reply_markup=humanizer_menu_kb(tone, mode, business))


@router.callback_query(F.data == "humanizer_menu")
async def cb_humanizer_menu(callback: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    tone     = data.get("tone", "College")
    mode     = data.get("mode", "Medium")
    business = data.get("business", False)
    text = await _humanizer_text(callback.from_user.id, tone, mode, business)
    await callback.message.edit_text(text, reply_markup=humanizer_menu_kb(tone, mode, business))
    await callback.answer()


# ── Настройки ─────────────────────────────────────────────────────────────────

@router.callback_query(F.data == "set_tone")
async def set_tone_menu(callback: CallbackQuery):
    await callback.message.edit_text("🎓 Выбери тон письма:", reply_markup=tone_kb())
    await callback.answer()


@router.callback_query(F.data.startswith("tone:"))
async def save_tone(callback: CallbackQuery, state: FSMContext):
    key = callback.data.split(":")[1]
    await state.update_data(tone=key)
    await callback.answer(f"✅ Тон: {TONE_OPTIONS.get(key, key)}")
    await cb_humanizer_menu(callback, state)


@router.callback_query(F.data == "set_mode")
async def set_mode_menu(callback: CallbackQuery):
    await callback.message.edit_text("📊 Выбери режим:", reply_markup=mode_kb())
    await callback.answer()


@router.callback_query(F.data.startswith("mode:"))
async def save_mode(callback: CallbackQuery, state: FSMContext):
    key = callback.data.split(":")[1]
    await state.update_data(mode=key)
    await callback.answer(f"✅ Режим: {MODE_OPTIONS.get(key, key)}")
    await cb_humanizer_menu(callback, state)


@router.callback_query(F.data == "toggle_business")
async def toggle_business(callback: CallbackQuery, state: FSMContext):
    data = await state.get_data()
    current = data.get("business", False)
    await state.update_data(business=not current)
    await callback.answer(f"✅ Бизнес-мод {'включён' if not current else 'выключен'}")
    await cb_humanizer_menu(callback, state)


# ── Запуск хуманизации ────────────────────────────────────────────────────────

@router.callback_query(F.data == "humanize_start")
async def start_humanize(callback: CallbackQuery, state: FSMContext):
    await state.set_state(HumanizerFSM.waiting_text)
    await callback.message.answer(
        "📝 Отправь текст для хуманизации.\n\n"
        "• Напиши сообщение\n"
        "• Или прикрепи файл <b>.txt / .docx / .doc</b>\n\n"
        "⚠️ Минимум 10 слов, максимум 20 000 слов.\n"
        "Для качественного результата рекомендуется 1200+ слов.",
        reply_markup=cancel_kb(),
    )
    await callback.answer()


# ── Обработка текста ──────────────────────────────────────────────────────────

async def _process_text(message: Message, state: FSMContext, text: str):
    text = text.strip()
    word_count = len(text.split())

    if word_count < 10:
        await message.answer("⚠️ Текст слишком короткий. Нужно минимум 10 слов.")
        return

    if word_count > 20000:
        await message.answer(
            f"⚠️ Текст слишком длинный ({word_count} слов). Максимум 20 000.\n"
            "Раздели на части."
        )
        return

    data = await state.get_data()
    tone     = data.get("tone", "College")
    mode     = data.get("mode", "Medium")
    business = data.get("business", False)

    cost = calculate_humanizer_cost(text, business=business)
    is_wl = settings.is_admin(message.from_user.id) or await db.is_whitelisted(message.from_user.id)

    if not is_wl:
        balance = await db.get_token_balance(message.from_user.id)
        if balance < cost:
            packages = settings.get_humanizer_packages()
            await message.answer(
                f"❌ <b>Недостаточно токенов!</b>\n\n"
                f"Нужно: <b>{cost:.1f} 🪙</b> ({word_count} слов)\n"
                f"Баланс: <b>{balance:.1f} 🪙</b>\n\n"
                f"Пополни баланс:",
                reply_markup=token_packages_kb(packages),
            )
            await state.clear()
            return

    chunks = math.ceil(word_count / 300)
    thinking_msg = await message.answer(
        f"⏳ Хуманизирую... ({word_count} слов"
        f"{', ' + str(chunks) + ' частей' if chunks > 1 else ''})\n"
        f"Ожидание ~{chunks * 70} сек."
    )

    try:
        result, words_spent = await humanize_text(text, tone=tone, mode=mode, business=business)

        if not is_wl:
            await db.deduct_tokens(message.from_user.id, cost)

        try:
            await thinking_msg.delete()
        except Exception:
            pass

        balance = await db.get_token_balance(message.from_user.id)
        header = (
            f"✅ <b>Готово!</b>\n"
            f"📊 Слов: {word_count} → {len(result.split())}\n"
            f"{'💰 Списано: <b>' + str(cost) + ' 🪙</b>' if not is_wl else '✨ Бесплатно (whitelist)'}\n"
            f"{'🪙 Баланс: <b>' + str(round(balance, 1)) + '</b>' if not is_wl else ''}\n"
            f"─────────────────────\n"
        )

        if len(header) + len(result) <= 4000:
            await message.answer(header, parse_mode="HTML")
            await message.answer(result, parse_mode=None,
                                 reply_markup=humanizer_menu_kb(tone, mode, business))
        else:
            await message.answer(header, parse_mode="HTML")
            file_bytes = result.encode("utf-8")
            doc = BufferedInputFile(file_bytes, filename="humanized.txt")
            await message.answer_document(
                doc,
                caption="📎 Результат (файл — текст слишком длинный для сообщения)",
                reply_markup=humanizer_menu_kb(tone, mode, business),
            )

    except Exception as e:
        logger.error(f"Humanizer error: {e}\n{traceback.format_exc()}")
        try:
            await thinking_msg.delete()
        except Exception:
            pass
        await message.answer(f"❌ Ошибка хуманайзера: {str(e)[:300]}")

    await state.clear()


@router.message(HumanizerFSM.waiting_text, F.text.in_(MENU_BUTTONS))
async def cancel_humanize(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Главное меню:", reply_markup=main_menu_kb())


@router.message(HumanizerFSM.waiting_text, F.text)
async def process_text_msg(message: Message, state: FSMContext):
    if len(message.text) >= 3500:
        await message.answer(
            f"⚠️ Сообщение слишком длинное — Telegram может его обрезать.\n"
            "Отправь текст файлом <b>.txt</b>",
        )
        return
    await _process_text(message, state, message.text)


@router.message(HumanizerFSM.waiting_text, F.document)
async def process_file(message: Message, state: FSMContext):
    doc: Document = message.document
    fname = (doc.file_name or "").lower()

    if not any(fname.endswith(ext) for ext in (".txt", ".docx", ".doc")):
        await message.answer("⚠️ Поддерживаются: <b>.txt, .docx, .doc</b>")
        return

    if doc.file_size > 500_000:
        await message.answer("⚠️ Файл слишком большой. Максимум 500 КБ.")
        return

    file = await message.bot.get_file(doc.file_id)
    downloaded = await message.bot.download_file(file.file_path)
    file_bytes = downloaded.read()

    try:
        if fname.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        elif fname.endswith(".docx"):
            import io
            from docx import Document as DocxDocument
            d = DocxDocument(io.BytesIO(file_bytes))
            text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        else:  # .doc
            try:
                import io
                from docx import Document as DocxDocument
                d = DocxDocument(io.BytesIO(file_bytes))
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
            except Exception:
                import re
                text = file_bytes.decode("utf-8", errors="ignore")
                text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', text)
                text = re.sub(r' +', ' ', text).strip()
    except Exception as e:
        await message.answer(f"❌ Не удалось прочитать файл: {e}")
        return

    if not text or len(text.strip()) < 20:
        await message.answer("❌ Не удалось извлечь текст. Попробуй сохранить как .txt")
        return

    await _process_text(message, state, text)
